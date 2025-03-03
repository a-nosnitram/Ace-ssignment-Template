import os
from docx import Document
from pathlib import Path
from string import Template
from colorama import Fore, Back, Style
from datetime import datetime

# report template
MD_REPORT_TEMPLATE = Template("""<div align="center">

# ${module_code} W${week_number} Assignment ${report_type}
## ${title}
### Tutor: ${tutor_name}
### Student: ${matriculation}
### Date: ${date}
</div>

---

## **Introduction**
Provide a brief overview of the topic, objectives, and scope of the assignment.

---

## **Design**
Explain your approach, methodology, and design decisions.

---

## **Testing**

### **Implementation Details**
Describe the development and coding process.

### **Testing Strategies**
Explain how you tested your program.

### **Test Results**
| **Test Description**   | **Expected Result**   | **Actual Result**   |
|------------------------|----------------------|----------------------|
| Test 1:                |                      |                      |
| Test 2:                |                      |                      |
| Test 3:                |                      |                      |


---

## **Conclusion**
Summarise key points, pretend you genuinely enjoyed the process of creating this, and talk about potential imporvements you'll never implemnent.

---

### **References** *(if needed)*
Include citations in academic format.
""")


def create_project():
    # get user input
    print(Fore.LIGHTBLUE_EX + "Welcome to the ultimate CS assignemnt template ! Enter your coursework details below" + Style.RESET_ALL)
    project_type = input_loop(
        "project type (Group (g) / Individual (i))", {"g", "i"})
    week_number = input_loop("week number (optional)", {"any"})
    title = input_loop("project title (optional)", {"any"})
    print(Fore.LIGHTBLACK_EX +
          "\nThis template creator provides basic project setups for the following languages :")
    print(Fore.BLUE + "Java, Pyhton, C, C++, Haskell, and JavaScript" +
          Fore.LIGHTBLACK_EX)
    print("Feel free to skip that part of the setup by pressing ENTER." + Style.RESET_ALL)
    language = input_loop("programming language", {"java", "python", "c",
                                                   "c++", "haskell",
                                                   "javascript", ""})

    # report details input
    module_code = input_loop("module code (optional)", {"any"})
    tutor_name = input_loop("tutor name (optional)", {"any"})
    matriculation = input_loop("matriculation number (oprional)", {"any"})
    report_format = input_loop(
        "report format (Markdown (m) / Word (w))", {"w", "m"})

    # folder name and base path based on input and folder name
    folder_name = f"W{week_number}-Assignment"
    base_path = Path(folder_name)

    if report_format == "m":
        report_extention = "md"
    else:
        report_extention = "docx"

    # create folder structure
    if project_type == "i":  # individual report setup
        src_path = base_path / "src"
        report_file = base_path / f"W{week_number}-Report." + report_extention

        report_type = ""
    else:  # group project setup
        src_path = base_path / "Code" / "src"
        report_path = base_path / "Reports"
        report_path.mkdir(parents=True, exist_ok=True)
        report_type_g = "Group Report"

        group_report_file = report_path / "Group-Report." + report_extention
        individual_report_file = report_path / "Individual-Report." + report_extention
        report_type_i = "Individual Report"

    # programming language specifics
    match language:
        case "java":
            main_file = src_path / "Main.java"
            main_file_content = """public class Main {
    public static void main(String args[]) {
        System.out.println("Good luck with your coursework ;)");
    }
}"""
            script_content = """#!/bin/bash
# this is a scrip to run your java project
# add files to the compile and run sections as needed

# compile your files
javac main.java

# run your files
java main"""
            script_file = src_path / "runJava.sh"
        case "c":
            main_file = src_path / "Main.c"
            main_file_content = """#include <stdio.h>

int main(int argc, char **argv) {
    printf("Good luck with your coursework ;)\n");
    return 0;
}"""
            script_content = """CC = GCC
TARGET = main
FLAGS = -Wall -Wextra
OBJS = main.o

$(TARGET): $(OBJS)
    $(CC) $(CFLAGS) -o $(TARGET) $(OBJS)

main.o: main.c
    $(CC) $(CFLAGS) -c main.c

clean:
    rm -f *.o $(TARGET)"""

            script_file = src_path / "Makefile"
        case "haskell":
            main_file = src_path / "Main.hs"
            main_file_content = """main :: IO ()
main = putStrLn "Good luck with your coursework ;)"
            """
            script_file = src_path / "runHaskell.sh"
            script_content = """#!/bin/bash

# this is a scrip to run your Haskell project
# add files to the compile and run sections as needed

# compile yout files
ghc -o main Main.hs

# run your files
./main
            """
        case "javscript":
            main_file = src_path / "main.js"
            main_file_content = """console.log("Good luck with your coursework ;)");"""
            script_file = src_path / "runJS.sh"
            script_content = """#!/bin/bash
# this is a scrip to run your JavaScript project
# add files as needed

# run your files
node main.js
"""
        case "python":
            main_file = src_path / "main.py"
            main_file_content = """if __name__ == "__main__":
    print("Good luck with your coursework ;)")"""
            script_file = src_path / "run_python.sh"
            script_content = """#!/bin/bash
# this is a scrip to run your JavaScript project
# add files as needed

# run your files
python3 main.py
"""
        case "c++":
            main_file = src_path / "main.cpp"
            main_file_content = """#include <iosteram>
int main() {
    std::cout << "Good luck with your coursework ;)" << std::endl;
    return 0;
}
"""
            script_file = src_path / "Makefile"
            script_content = """CXX=g++
TARGET=main
CXXFLAGS = -Wall -Wextra
OBJS = main.o

$(TARGET): $(OBJS)
    $(CXX) $(CXXFLAGS) -o $(TARGET) $(OBJS)

main.o: main.cpp
    $(CXX) $(CXXFLAGS) -c main.cpp

clean:
    rm -f *.o $(TARGET)
"""

    # create directories using Path.mkdir
    src_path.mkdir(parents=True, exist_ok=True)

    # configure project files
    if language != "":
        main_file.write_text(main_file_content)
        script_file.write_text(script_content)

    date={datetime.today().strftime('%d %m %Y')}
    make_report(report_extention, report_type, week_number, module_code, tutor_name, title, matriculation, date)

    print(Fore.GREEN + "All done :з project created !" + Style.RESET_ALL)


def input_loop(prompt, target_set):
    inp = input(f"Enter {prompt} : ").strip().lower()

    if "any" in target_set:
        return inp

    while inp not in target_set:
        print(Fore.RED + "Invalid input" + Style.RESET_ALL)
        inp = input(f"Enter {prompt} : ").strip().lower()
    return inp


def make_report(report_extention, project_type):
    if report_extention == "w":
        make_word_report(project_type)
    else:
        make_md_report(project_type)


def make_word_report(project_type):
    if report_type == "i":
        doc = make_word("")
        doc.save(f"W{week_number}-Report.docx")
    else: 
        doc_ind = make_word("Individual Report")
        doc_gr = make_word("Group Report")
        doc_ind.save("Individual-Report.docx")
        doc_gr.save("Group-Report.docx")


def make_word(rep_t):
        doc = Document()

        doc.add_heading(
            f"{module_code} W{week_number} Assignment {rep_t}", level=1)
        doc.add_paragraph(
            f"Title: {title}\nTutor: {tutor_name}\nStudent: {matriculation}\nDate: {date}\n")

        doc.add_paragraph("## **Introduction**")
        doc.add_paragraph(
            "Provide a brief overview of the topic, objectives, and scope of the assignment.")

        doc.add_paragraph("## **Design**")
        doc.add_paragraph(
            "Explain your approach, methodology, and design decisions.")

        doc.add_paragraph("## **Testing**")
        doc.add_paragraph("### **Implementation Details**")
        doc.add_paragraph("Describe the development and coding process.")

        doc.add_paragraph("### **Testing Strategies**")
        doc.add_paragraph("Explain how you tested your program.")

        doc.add_paragraph("### **Test Results**")

        table = doc.add_table(rows=4, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Test Description"
        hdr_cells[1].text = "Expected Result"
        hdr_cells[2].text = "Actual Result"

        for i in range(1, 4):
            row_cells = table.rows[i].cells
            row_cells[0].text = f"Test {i}"
            row_cells[1].text = ""
            row_cells[2].text = ""

        doc.add_paragraph("## **Conclusion**")
        doc.add_paragraph(
            "Summarize key points, pretend you genuinely enjoyed the process of creating this, and talk about potential improvements you'll never implement.")

        doc.add_paragraph("### **References** (if needed)")
        doc.add_paragraph("Include citations in academic format.")

        return doc


def make_md_report(project_type):
    # define report content using substitute
    if project_type == "i":
        report_contents = MD_REPORT_TEMPLATE.substitute(module_code=module_code, week_number=week_number, report_type=report_type, tutor_name=tutor_name, title=title, matriculation=matriculation, date=date)
        # write report files
        report_file.write_text(report_contents)
    else:
        report_g_contents = MD_REPORT_TEMPLATE.substitute(module_code=module_code, week_number=week_number, report_type=report_type_g, tutor_name=tutor_name, title=title, matriculation=matriculation, date=date)
        group_report_file.write_text(report_g_contents)

        report_i_contents = MD_REPORT_TEMPLATE.substitute(module_code=module_code, week_number=week_number, report_type=report_type_i, tutor_name=tutor_name, title=title, matriculation=matriculation, date=date)
        individual_report_file.write_text(report_i_contents)

if __name__ == "__main__":
    create_project()
